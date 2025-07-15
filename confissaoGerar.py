import os
import re
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pdf2image import convert_from_path
from collections import defaultdict
from censurar import aplicar_censura, censurar_cpf

class GeradorDocumentos:
    def __init__(self, pasta_txts='txts', pasta_destino='docxs_gerados', planilha_resultado='dados_atualizados.xlsx', censurar_dados=False):
        self.pasta_txts = pasta_txts
        self.pasta_destino = pasta_destino
        self.df = pd.read_excel(planilha_resultado)
        os.makedirs(self.pasta_destino, exist_ok=True)
        self.contador_nomes = defaultdict(int)
        self.censurar_dados = censurar_dados

    def _inserir_pdf_como_imagem(self, doc, caminho_pdf): # Insere cada página de um PDF como imagem no documento.
        try:
            paginas = convert_from_path(caminho_pdf, dpi=200, poppler_path=r'C:\Poppler\Library\bin')
        except Exception as e:
            print(f"Erro ao converter PDF {caminho_pdf}: {e}")
            print("Verifique se o Poppler está instalado e o caminho está correto.")
            return

        for imagem in paginas:
            caminho_temp = "temp_pag.pdf.png"

            # Cortar metade inferior
            imagem_pil = imagem.copy()
            largura_px, altura_px = imagem_pil.size
            imagem_cortada = imagem_pil.crop((0, 0, largura_px, altura_px // 2))
            imagem_cortada.save(caminho_temp, 'PNG')

            # Proporção da imagem
            proporcao = imagem_cortada.width / imagem_cortada.height

            # Limites da página A4 (em cm)
            largura_max_cm = 16  # margem segura (21cm - 2.5cm cada lado)
            altura_max_cm = 24  # margem segura (29.7cm - 2.85cm cada lado)

            # Tentativa 1: usar largura máxima
            largura_final_cm = largura_max_cm
            altura_final_cm = largura_final_cm / proporcao

            # Se ultrapassar a altura da página, reajustar
            if altura_final_cm > altura_max_cm:
                altura_final_cm = altura_max_cm
                largura_final_cm = altura_final_cm * proporcao

            # Inserir no documento
            doc.add_page_break()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(caminho_temp, width=Cm(largura_final_cm), height=Cm(altura_final_cm))
        if os.path.exists(caminho_temp):
            os.remove(caminho_temp)


    def _numero_por_extenso(self, valor): # Converte um número (float) em sua representação por extenso em português.
        unidades = ["", "um", "dois", "três", "quatro", "cinco", "seis", "sete", "oito", "nove"]
        dezenas = ["", "dez", "vinte", "trinta", "quarenta", "cinquenta", "sessenta", "setenta", "oitenta", "noventa"]
        especiais = ["dez", "onze", "doze", "treze", "quatorze", "quinze", "dezesseis", "dezessete", "dezoito", "dezenove"]
        centenas = ["", "cem", "duzentos", "trezentos", "quatrocentos", "quinhentos", "seiscentos", "setecentos", "oitocentos", "novecentos"]

        valor = float(valor.replace('.', '').replace(',', '.'))

        def extenso_menor_1000(n):
            if n == 0:
                return ""
            partes = []
            if n >= 100:
                c = n // 100
                if c == 1 and n % 100 == 0:
                    partes.append("cem")
                else:
                    partes.append(centenas[c])
                n = n % 100
            if 10 <= n <= 19:
                partes.append(especiais[n - 10])
            else:
                d = n // 10
                if d:
                    partes.append(dezenas[d])
                u = n % 10
                if u:
                    partes.append(unidades[u])
            return " e ".join([p for p in partes if p])

        inteiro = int(valor)
        centavos = round((valor - inteiro) * 100)

        partes_texto = []

        if inteiro == 0:
            partes_texto.append("zero")
        elif inteiro == 1:
            partes_texto.append("um real")
        elif inteiro > 1:
            if inteiro >= 1000:
                milhar = inteiro // 1000
                resto_milhar = inteiro % 1000
                if milhar == 1:
                    partes_texto.append("mil")
                else:
                    partes_texto.append(f"{extenso_menor_1000(milhar)} mil")
                if resto_milhar > 0:
                    partes_texto.append(extenso_menor_1000(resto_milhar))
            else:
                partes_texto.append(extenso_menor_1000(inteiro))

            if inteiro == 1:
                partes_texto[-1] += " real"
            elif inteiro > 1:
                if not partes_texto[-1].endswith("mil"): # Evita "mil reais"
                    partes_texto[-1] += " reais"

        resultado = " e ".join(partes_texto).replace("mil reais", "mil") # Ajuste para "mil reais"

        if centavos > 0:
            if inteiro > 0:
                resultado += " e "
            if centavos == 1:
                resultado += "um centavo"
            else:
                resultado += f"{extenso_menor_1000(centavos)} centavos"

        return resultado.strip()


    def _criar_paragrafo_borda(self, doc, texto):
        """Cria um parágrafo com bordas e formatação específica."""
        p = doc.add_paragraph()
        run = p.add_run(texto)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Garamond'
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5

        p_border = OxmlElement('w:pBdr')
        for border_name in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '1')
            border.set(qn('w:color'), 'auto')
            p_border.append(border)
        p._p.get_or_add_pPr().append(p_border)
        return p

    def _extrair_cidade_txt(self, caminho_txt):
        """Extrai a cidade de um arquivo de texto."""
        with open(caminho_txt, 'r', encoding='utf-8') as f:
            texto = f.read().lower()

        municipios = [m.start() for m in re.finditer('municipio de ', texto)]
        if len(municipios) < 2:
            return "SANTA CRUZ DO SUL"

        inicio = municipios[1] + len('município de ')
        fim = texto.find(',', inicio)
        cidade = texto[inicio:fim].strip().upper()
        return cidade

    def _extrair_texto_parcial(self, caminho_txt):
        """Extrai um trecho específico de texto de um arquivo."""
        with open(caminho_txt, 'r', encoding='utf-8') as f:
            texto = f.read()

        inicio = texto.lower().find('outro lado,')
        if inicio == -1:
            return "[Texto não encontrado no .txt]"

        inicio_real = inicio + len('outro lado,')
        trecho = texto[inicio_real:].strip()
        contador_virgulas = 0
        fim = 0
        for i, c in enumerate(trecho):
            if c == ',':
                contador_virgulas += 1
            if contador_virgulas > 10:
                fim = i
                break

        return trecho[:fim].strip()

    def _aplicar_censura_condicional(self, texto_buffer, is_name=False):
        """
        Aplica censura a um texto ou CPF se a censura estiver ativada.
        Args:
            texto_buffer (str): O texto a ser processado.
            is_name (bool): True se o texto_buffer for um nome a ser censurado.
        Returns:
            str: O texto censurado ou original.
        """
        if self.censurar_dados:
            if is_name:
                return aplicar_censura(texto_buffer, self.censurar_dados)
            else:
                regex_cpf = re.compile(r'\d{3}[.\s]?\d{3}[.\s]?\d{3}[-\s]?\d{2}')
                cpf_encontrado = regex_cpf.search(texto_buffer)
                if cpf_encontrado:
                    cpf_original = cpf_encontrado.group()
                    cpf_censurado = censurar_cpf(cpf_original, self.censurar_dados)
                    return texto_buffer.replace(cpf_original, cpf_censurado)
                return texto_buffer
        return texto_buffer

    def gerar_documentos(self):
        """Gera os documentos Word com base nos dados e arquivos de texto."""
        for arquivo in os.listdir(self.pasta_txts):
            if arquivo.endswith('.txt'):
                nome_base = os.path.splitext(arquivo)[0]

                if '_' in nome_base and nome_base.split('_')[-1].isdigit():
                    nome_base_corrigido = '_'.join(nome_base.split('_')[:-1])
                else:
                    nome_base_corrigido = nome_base

                info = self.df[self.df['Arquivo'].str.contains(nome_base_corrigido, na=False, regex=False)]

                if info.empty:
                    print(f"Dados não encontrados na planilha para o arquivo: {arquivo}")
                    continue

                nome_devedor = info.iloc[0]['Nome']
                valor_divida = info.iloc[0]['Dívida Atualizada']

                doc = Document()

                # Cabeçalho
                section = doc.sections[0]
                header = section.header
                header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                run_header = header_paragraph.add_run()
                run_header.add_picture('cab.png', width=Pt(500))

                # Rodapé
                footer = section.footer
                footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                run_footer = footer_paragraph.add_run()
                run_footer.add_picture('rod.png', width=Pt(500))

                cidade_extraida = self._extrair_cidade_txt(os.path.join(self.pasta_txts, arquivo))
                texto_borda = f"EXCELENTÍSSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) DE DIREITO DA VARA JUDICIAL DA COMARCA DE {cidade_extraida}, RS."
                self._criar_paragrafo_borda(doc, texto_borda)

                p = doc.add_paragraph()
                partes = [
                    ("AFUBRA – ASSOCIAÇÃO DOS FUMICULTORES DO BRASIL,", True, False),
                    (" pessoa jurídica de direito privado, inscrita no CNPJ/MF sob o nº 95.430.690/0001-25, com sede à Rua Júlio de Castilhos, nº 1030, centro, no Município de Santa Cruz do Sul, no Estado do Rio Grande do Sul, vem, à presença de Vossa Excelência, por sua procuradora firmatária, qualificada conforme procuração anexa, que recebe intimações, apenas e tão somente, ", False, False),
                    ("sob pena de nulidade, em nome de Cleidimara da Silva Flores,", True, True),
                    (" brasileira, casada, advogada, inscrita na OAB/RS sob o nº 63.984 e CPF/MF nº 983.244.380-68, endereço eletrônico adv@floresadv.com.br, com escritório profissional na Rua Carlos Trein Filho, nº 600, em Santa Cruz do Sul, RS, propor a presente", False, False)
                ]
                for texto, negrito, sublinhado in partes:
                    run = p.add_run(texto)
                    run.bold = negrito
                    run.underline = sublinhado
                    run.font.size = Pt(12)
                    run.font.name = 'Garamond'
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.left_indent = Pt(108)

                p = self._criar_paragrafo_borda(doc, "AÇÃO DE EXECUÇÃO POR QUANTIA CERTA, com fundamento nos artigos 784, III e 824 e seguintes do Código de Processo Civil, contra")
                p.paragraph_format.left_indent = Pt(108)

                trecho_txt = self._extrair_texto_parcial(os.path.join(self.pasta_txts, arquivo))
                p = doc.add_paragraph()
                partes_paragrafo_devedor = []

                buffer = ''
                virgula_count = 0

                for c in trecho_txt:
                    if c == ',':
                        texto_buffer = buffer.strip()

                        if virgula_count == 0: # É o nome
                            nome_tratado = self._aplicar_censura_condicional(texto_buffer, is_name=True)
                            partes_paragrafo_devedor.append((nome_tratado.upper(), True))
                        else:
                            texto_processado = self._aplicar_censura_condicional(texto_buffer)
                            partes_paragrafo_devedor.append((texto_processado, False))

                        buffer = ''
                        virgula_count += 1
                    else:
                        buffer += c

                # Última parte (fora do loop)
                if buffer:
                    texto_processado = self._aplicar_censura_condicional(buffer.strip())
                    partes_paragrafo_devedor.append((texto_processado, False))

                # Geração do parágrafo do devedor
                for i, (texto, negrito) in enumerate(partes_paragrafo_devedor):
                    run = p.add_run(texto)
                    run.bold = negrito
                    run.font.size = Pt(12)
                    run.font.name = 'Garamond'
                    if i < len(partes_paragrafo_devedor) - 1 and texto != "":
                        p.add_run(', ').font.size = Pt(12)

                p.add_run(", RS, pelos fatos e fundamentos que seguem:").font.size = Pt(12)

                for run in p.runs:
                    run.font.name = 'Garamond'

                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.left_indent = Pt(108)

                doc.add_page_break()

                # Parte dos fatos
                self._criar_paragrafo_borda(doc, "I – DOS FATOS")

                for texto in [
                    "O(A) Executado(a) contraiu dívida com a Exequente conforme consta no Instrumento Particular de Confissão de Dívida anexo.",
                    "Ocorre que passada a data de vencimento fixada entre as partes para o adimplemento do valor devido, não houve o pagamento da totalidade do débito.",
                    "Portanto, não restou alternativa por parte da Exequente, senão promover a presente Execução."
                ]:
                    paragrafo_fatos = doc.add_paragraph(texto)
                    paragrafo_fatos.paragraph_format.first_line_indent = Pt(108)
                    paragrafo_fatos.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    paragrafo_fatos.paragraph_format.line_spacing = 1.5
                    for run_fatos in paragrafo_fatos.runs:
                        run_fatos.font.size = Pt(12)
                        run_fatos.font.name = 'Garamond'
                doc.add_paragraph("")

                # Parte do direito
                self._criar_paragrafo_borda(doc, "II – DO DIREITO")
                textos_direito = [
                    ("Conforme preconiza o Código de Processo Civil acerca dos requisitos básicos para o ajuizamento de qualquer ação de execução, ante ao inadimplemento de devedor, vem o artigo 786 assim dispor:", 12, True),
                    ("Art. 786 – A execução pode ser instaurada caso o devedor não satisfaça a obrigação certa, líquida e exigível, consubstanciada em título executivo.", 10, False),
                    ("", 12, True),
                    ("Ademais, é considerado título executivo extrajudicial o documento particular assinado pelo devedor e por 2 (duas) testemunhas, consoante preconiza o art. 784, III, do CPC.", 12, True),
                    ("Portanto, plenamente cabível o ajuizamento da presente Ação de Execução, uma vez que acompanhada de título executivo extrajudicial, qual seja: Termo de Confissão de Dívida, assinado pelo devedor e duas testemunhas, consistindo em pagamento de valor certo, líquido e exigível.", 12, True),
                    ("O direito da ação relativo ao Contrato de Confissão de Dívida no Código Civil, prescreve em cinco (05) anos, conforme se verifica do artigo 206, § 5º, I, in verbis:", 12, True),
                    ("Art. 206. Prescreve:", 10, False),
                    ("(...)" , 10, False),
                    ("§ 5º Em cinco anos:", 10, False),
                    ("(...)" , 10, False),
                    ("I - a pretensão de cobrança de dívidas líquidas constantes de instrumento público ou particular;", 10, False),
                    ("", 12, True),
                    ("Tal título possui, portanto, eficácia de título executivo, possibilitando o ingresso da ação de execução por quantia certa, espécie de execução que tem por objeto expropriar bens do Executado, a fim de satisfazer o direito da Exequente (art. 824 do CPC).", 12, True),
                    ("Assim, existindo 'legitimatio ad causam', interesse processual, e sendo o pedido juridicamente possível, encontra-se apto para a prestação da tutela jurisdicional que adiante se invocará.", 12, True)
                ]
                for texto, tamanho, espacamento in textos_direito:
                    if "in verbis" in texto or "legitimatio ad causam" in texto:
                        p_direito = doc.add_paragraph()
                        # Usa regex para encontrar o termo e dividir
                        match = re.search(r'(in verbis|legitimatio ad causam)', texto)
                        if match:
                            termo = match.group(0)
                            partes = texto.split(termo)
                            p_direito.add_run(partes[0]).font.size = Pt(tamanho)
                            run = p_direito.add_run(termo)
                            run.italic = True
                            run.font.size = Pt(tamanho)
                            run.font.name = 'Garamond'
                            if len(partes) > 1:
                                p_direito.add_run(partes[1]).font.size = Pt(tamanho)
                        else: # Fallback se não encontrar os termos esperados
                            p_direito.add_run(texto).font.size = Pt(tamanho)
                    else:
                        p_direito = doc.add_paragraph(texto)

                    if tamanho == 10:
                        p_direito.paragraph_format.left_indent = Pt(108)
                        p_direito.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_direito.paragraph_format.space_after = Pt(0)
                        p_direito.paragraph_format.space_before = Pt(0)
                        p_direito.paragraph_format.line_spacing = 1.0
                    else:
                        p_direito.paragraph_format.first_line_indent = Pt(108)
                        p_direito.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_direito.paragraph_format.line_spacing = 1.5
                    for run_direito in p_direito.runs:
                        run_direito.font.size = Pt(tamanho)
                        run_direito.font.name = 'Garamond'

                # Parte dos pedidos
                self._criar_paragrafo_borda(doc, "III – DOS PEDIDOS")
                extenso_divida = self._numero_por_extenso(valor_divida)

                p = doc.add_paragraph()
                run = p.add_run("ANTE O EXPOSTO")
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = 'Garamond'

                run = p.add_run(", requer a Vossa Excelência:")
                run.font.size = Pt(12)
                run.font.name = 'Garamond'

                p.paragraph_format.first_line_indent = Pt(108)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.5

                nome_devedor_tratado = self._aplicar_censura_condicional(nome_devedor, is_name=True)

                pedidos = [
                    ("Que receba a presente Ação de Execução por Quantia certa, determinando, desde já, a expedição de Mandado de Citação do(a) Executado(a), para, no prazo de 03 (três) dias, satisfazer a Exequente, pagando o valor de",
                     valor_divida, extenso_divida,
                     " referente a dívida atualizada até 23/04/2025, acrescendo-se a atualização monetária pelo IGP-M, juros moratórios de 1 % ao mês, multa de 10%, mais custas processuais e honorários advocatícios, até a data do efetivo pagamento, sob pena de serem penhorados tantos bens quantos necessários para garantir a dívida; "),
                    ("Vencido este prazo e não efetuado o pagamento, proceda de imediato à penhora de ativos financeiros, através do Sistema SISBAJUD, localizados em nome do(a) Executado(a) ",
                     None, None, f" {nome_devedor_tratado}."),
                    ("Caso não sejam localizados ativos financeiros em nome do(a) Executado(a), que o Oficial de Justiça, munido com a segunda via do mandado, proceda de imediato a penhora e avaliação de tantos bens quanto bastem para garantia do crédito reclamado, mais juros, custas e honorários;",
                     None, None, ""),
                    ("Em caso de penhora de veículos, em atenção ao que dispõe o § 2º do art. 840, do CPC, desde já a Exequente esclarece que não concorda que o(s) bem(ns) seja(m) depositado(s) com o(a) Executado(a), pois a continuidade da utilização do(s) veículo(s) ocasionará a sua depreciação ou a perda, além da possibilidade da prática de infrações de trânsito, gerando despesas. Assim, requer que o(s) veículo(s) seja(m) depositado(s) com o depositário judicial, nos termos do inc. II, do art. 840, do CPC. Não havendo na Comarca depositário judicial, requer que seja(m) depositado(s) em poder desta Exequente, que se compromete com a sua conservação;",
                     None, None, ""),
                    ("Não sendo encontrado(a) o(a) Executado(a), ou em caso deste tentar furtar-se da presente execução, que lhe seja arrestado bens suficientes, independente de novo mandado, dando-se ciência à Exequente para as providências de citação editalícia, prevista no art. 830, § 2º do CPC;",
                     None, None, ""),
                    ("Após a efetivação da penhora, seja expedido Mandado de Averbação da penhora, no registro competente;",
                     None, None, ""),
                    ("Sejam fixados os honorários advocatícios em 20% (vinte por cento) sobre o valor atualizado da causa, nos termos do art. 85, § 2º do CPC;",
                     None, None, ""),
                    ("Seja expedida a certidão de existência da presente Execução para que a Exequente proceda a averbação no registro dos bens sujeitos à penhora, nos termos do art. 828, do CPC;",
                     None, None, ""),
                    ("Que as intimações sejam efetuadas apenas e tão somente, sob pena de nulidade, em nome de Cleidimara da Silva Flores, com endereço profissional na Rua Carlos Trein Filho, nº 600, CEP 96.810-176, em Santa Cruz do Sul, RS.",
                     None, None, "")
                ]

                for i, (inicio_texto, valor, extenso, fim_texto) in enumerate(pedidos):
                    p_pedido = doc.add_paragraph(style='List Number')

                    if i == 3 and "não concorda" in inicio_texto:
                        partes = inicio_texto.split("não concorda")
                        p_pedido.add_run(partes[0]).font.size = Pt(12)
                        run_nc = p_pedido.add_run("não concorda")
                        run_nc.bold = True
                        run_nc.underline = True
                        run_nc.font.size = Pt(12)
                        run_nc.font.name = 'Garamond'
                        p_pedido.add_run(partes[1] + " ").font.size = Pt(12)

                    elif i == 7:
                        trecho_alvo = "expedida a certidão de existência da presente Execução para que a Exequente proceda a averbação no registro dos bens sujeitos à penhora"
                        if trecho_alvo in inicio_texto:
                            antes, depois = inicio_texto.split(trecho_alvo)
                            p_pedido.add_run(antes).font.size = Pt(12)

                            run_destacado = p_pedido.add_run(trecho_alvo)
                            run_destacado.bold = True
                            run_destacado.underline = True
                            run_destacado.font.size = Pt(12)
                            run_destacado.font.name = 'Garamond'

                            p_pedido.add_run(depois + " ").font.size = Pt(12)
                        else:
                            p_pedido.add_run(inicio_texto + " ").font.size = Pt(12)

                    elif i == 8:
                        texto = inicio_texto
                        parte1, resto = texto.split("sob pena de nulidade")
                        parte2, parte3 = resto.split("Cleidimara da Silva Flores")

                        p_pedido.add_run(parte1).font.size = Pt(12)

                        run_sub = p_pedido.add_run("sob pena de nulidade")
                        run_sub.underline = True
                        run_sub.font.size = Pt(12)
                        run_sub.font.name = 'Garamond'

                        p_pedido.add_run(parte2).font.size = Pt(12)

                        run_nome = p_pedido.add_run("Cleidimara da Silva Flores")
                        run_nome.bold = True
                        run_nome.font.size = Pt(12)
                        run_nome.font.name = 'Garamond'

                        p_pedido.add_run(parte3 + " ").font.size = Pt(12)

                    else:
                        p_pedido.add_run(inicio_texto + " ").font.size = Pt(12)

                    for run_pedido in p_pedido.runs:
                        run_pedido.font.name = 'Garamond'

                    if valor is not None and extenso is not None:
                        run_valor = p_pedido.add_run(f"R$ {valor} ")
                        run_valor.bold = True
                        run_valor.font.size = Pt(12)
                        run_valor.font.name = 'Garamond'

                        run_extenso = p_pedido.add_run(f"({extenso})")
                        run_extenso.bold = True
                        run_extenso.font.size = Pt(12)
                        run_extenso.font.name = 'Garamond'

                    if i == 1:
                        run_fim = p_pedido.add_run(fim_texto.upper())
                        run_fim.bold = True
                    else:
                        run_fim = p_pedido.add_run(fim_texto)

                    run_fim.font.size = Pt(12)
                    run_fim.font.name = 'Garamond'

                    p_pedido.paragraph_format.left_indent = Pt(144)
                    p_pedido.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p_pedido.paragraph_format.line_spacing = 1.5

                # Última linha com valor da causa e valor por extenso
                p_valor_causa = doc.add_paragraph()
                run = p_valor_causa.add_run("Dá-se à presente causa o valor de ")
                run.font.size = Pt(12)
                run.font.name = 'Garamond'

                run_valor = p_valor_causa.add_run(f"R$ {valor_divida} ")
                run_valor.bold = True
                run_valor.font.size = Pt(12)
                run_valor.font.name = 'Garamond'

                run_extenso = p_valor_causa.add_run(f"({extenso_divida})")
                run_extenso.bold = True
                run_extenso.font.size = Pt(12)
                run_extenso.font.name = 'Garamond'

                p_valor_causa.paragraph_format.first_line_indent = Pt(108)
                p_valor_causa.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_valor_causa.paragraph_format.line_spacing = 1.5

                # Assinatura e encerramento
                linhas_finais = [
                    ("Nestes termos,", 'left', False),
                    ("Pede deferimento.", 'left', False),
                    ("", 'left', False),
                    ("Santa Cruz do Sul, RS, 24 de abril de 2025.", 'center', False),
                    ("", 'center', False),
                    ("Cleidimara da Silva Flores", 'center', True),
                    ("OAB/RS 63.984", 'center', True)
                ]

                for texto, alinhamento, negrito in linhas_finais:
                    p_final = doc.add_paragraph(texto)

                    if alinhamento == 'left':
                        p_final.paragraph_format.first_line_indent = Pt(108)
                    elif alinhamento == 'center':
                        p_final.paragraph_format.first_line_indent = Pt(0)
                        p_final.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    p_final.paragraph_format.line_spacing = 0.7

                    for run_final in p_final.runs:
                        run_final.font.size = Pt(12)
                        run_final.font.name = 'Garamond'
                        run_final.bold = negrito

                # Salva o documento
                self.contador_nomes[nome_base] += 1
                contador = self.contador_nomes[nome_base]

                if contador == 1:
                    nome_final = f"{nome_base}.docx"
                else:
                    nome_final = f"{nome_base}_{contador}.docx"

                nome_pdf = nome_final.replace('.docx', '.pdf')
                if nome_pdf.endswith('_1.pdf'):
                    nome_pdf = nome_pdf.replace('_1.pdf', '.pdf')

                caminho_pdf = os.path.join('pdfs_formatados', nome_pdf)

                if os.path.exists(caminho_pdf):
                    self._inserir_pdf_como_imagem(doc, caminho_pdf)
                else:
                    print(f"⚠️ PDF correspondente não encontrado para {nome_final}")

                caminho_destino = os.path.join(self.pasta_destino, nome_final)
                doc.save(caminho_destino)

        print("✅ Documentos gerados com sucesso!")

# Função para ser chamada externamente, controlando a criação da instância da classe
def executar(censurar: bool):
    """
    Inicializa e executa o processo de geração de documentos.
    Args:
        censurar (bool): Se True, aplica a censura de dados (nomes e CPFs).
    """
    gerador = GeradorDocumentos(censurar_dados=censurar)
    gerador.gerar_documentos()